# task_3_final.py
import os
import re
import json
import streamlit as st
from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader
import google.generativeai as gen

# ------------------ CONFIG ------------------
GEN_API_KEY = "AIzaSyBr6nZMQnuGHSC9Ynzgv90SWt9Hebz8KD0"
gen.configure(api_key=GEN_API_KEY)

OUTPUT_DIR = "task_3_output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

MODEL_NAME = "gemini-2.5-flash"

# ------------------ HELPERS ------------------
def extract_text_from_pdfs(files):
    """Extract text from multiple PDFs and concatenate."""
    text = ""
    for f in files:
        reader = PdfReader(f)
        for page in reader.pages:
            content = page.extract_text()
            if content:
                text += content + "\n"
    return text

def extract_placeholders(template):
    """Get placeholders from the template."""
    doc = Document(template)
    collected = ""
    for p in doc.paragraphs:
        collected += p.text + "\n"
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                collected += cell.text + "\n"
    keys = set()
    keys.update(re.findall(r"\[(.*?)\]", collected))
    keys.update(re.findall(r"\{\{(.*?)\}\}", collected))
    return [k.strip() for k in keys]

def normalize_key(key):
    """Normalize keys to UPPERCASE_WITH_UNDERSCORES"""
    key = key.upper().strip()
    key = re.sub(r"[^A-Za-z0-9]+", "_", key)
    return key

# ------------------ PRE-EXTRACTION ------------------
def pre_extract_fields(report_text):
    """Extract key fields using regex as a backup to LLM."""
    result = {
        "INSURED_H_STREET": "",
        "INSURED_H_CITY": "",
        "INSURED_H_STATE": "",
        "INSURED_H_ZIP": "",
        "DATE_LOSS": "",
        "DATE_INSPECTED": "",
        "DATE_RECEIVED": "",
        "INSURED_NAME": "",
        "MORTGAGEE": "",
        "MORTGAGE_CO": "",
        "TOL_CODE": "",
    }
    try:
        match = re.search(r"Insured:\s*(.+)", report_text)
        if match:
            result["INSURED_NAME"] = match.group(1).strip()

        match = re.search(
            r"Risk address\s*\n(.+?)\n(.+?),\s*([A-Z]{2})\s*(\d{5})", report_text
        )
        if match:
            result["INSURED_H_STREET"] = match.group(1).strip()
            result["INSURED_H_CITY"] = match.group(2).strip()
            result["INSURED_H_STATE"] = match.group(3).strip()
            result["INSURED_H_ZIP"] = match.group(4).strip()

        match = re.search(r"Date Taken\s*[:\-]?\s*(\d{1,2}/\d{1,2}/\d{4})", report_text)
        if match:
            result["DATE_INSPECTED"] = match.group(1)

        match = re.search(
            r"mortgage company was verified as\s*([^\n.]+)", report_text, re.IGNORECASE
        )
        if match:
            result["MORTGAGE_CO"] = match.group(1).strip()

        match = re.search(r"Date of Loss\s*[:\-]?\s*(\d{1,2}/\d{1,2}/\d{4})", report_text)
        if match:
            result["DATE_LOSS"] = match.group(1)

    except Exception as e:
        print("Pre-extraction error:", e)
    return result

# ------------------ GEMINI LLM CALL ------------------
def call_llm(fields, full_report_text):
    model = gen.GenerativeModel(MODEL_NAME)
    field_template = {normalize_key(f): "" for f in fields}
    prompt = f"""
You are an insurance extraction AI.

Extract values and return ONLY valid JSON.

Rules:
- Do NOT change field names.
- Use exact wording from the report.
- Extract full addresses, city, state, zip even if split across lines.
- If a field is missing, return "".
- Dates must be copied exactly.

FIELDS:
{json.dumps(field_template, indent=2)}

REPORT:
\"\"\"{full_report_text}\"\"\"
"""
    try:
        response = model.generate_content(prompt)
        response_text = response.text.strip()
        response_text = response_text.replace("```json", "").replace("```", "")
        return json.loads(response_text)
    except Exception as e:
        return {}

# ------------------ TEMPLATE FILL ------------------
def fill_template(template, mapping):
    doc = Document(template)
    def replace(obj):
        for key, value in mapping.items():
            obj.text = obj.text.replace(f"[{key}]", value)
            obj.text = obj.text.replace(f"{{{{{key}}}}}", value)
    for p in doc.paragraphs:
        replace(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace(p)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ------------------ STREAMLIT UI ------------------
st.title("📄 GLR Insurance Automation Pipeline (Gemini Powered)")

template_file = st.file_uploader("Upload Template (.docx)", type=["docx"])
pdf_files = st.file_uploader("Upload PDF Reports (multiple allowed)", accept_multiple_files=True, type=["pdf"])

if st.button("Run"):
    if not template_file or not pdf_files:
        st.error("Upload both template and at least one PDF report.")
        st.stop()

    st.write("🔍 Extracting placeholders from template...")
    placeholders = extract_placeholders(template_file)
    st.write(placeholders)

    st.write("📑 Extracting text from PDFs...")
    report_text = extract_text_from_pdfs(pdf_files)
    st.text_area("PDF Text Preview:", report_text[:3000], height=200)

    st.write("🔧 Pre-extracting structured fields...")
    pre_fields = pre_extract_fields(report_text)
    st.json(pre_fields)

    st.write("🤖 Calling Gemini LLM...")
    ai_fields = call_llm(placeholders, report_text)
    st.json(ai_fields)

    # Safe merge
    final_mapping = {**ai_fields, **pre_fields}

    st.success("📌 Fields merged successfully.")
    st.json(final_mapping)

    st.write("📄 Generating final GLR report...")
    final_doc = fill_template(template_file, final_mapping)

    st.download_button(
        "⬇ Download Completed GLR Report",
        data=final_doc,
        file_name="Final_GLR_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    st.success("🎉 GLR Report generated successfully!")
